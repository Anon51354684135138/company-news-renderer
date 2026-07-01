""".docx painting backend for the PubCo Brief — a faithful "cook" that receives
the recipe from `brief_model.build_blocks()` and paints it into a Word document.

Architecture (2026-06-19 recipe+cooks refactor):
- `brief_model` is the structure + rules authority: it builds the ordered list
  of typed block records ("the recipe"), runs the content rules (star-landing,
  source backstop, ★ stamping), and owns `SectionItem`, the guard helpers, and
  all nine block types.
- This module owns ONLY the Word painting surface: Times New Roman, strictly
  black text (links underlined black), 0.5" margins, confidentiality footer with
  a PAGE field, and the red ★ as the sole permitted color accent.
- `render_docx(blocks)` is the core painter; `render_brief`/`save_brief` are the
  frozen public wrappers (= build_blocks(**content) → render_docx). They are
  frozen to protect the live cloud routines' rendering call.
- Dependency direction: this module imports from `brief_model`; it does NOT
  export helpers to the other format backends (brief_pdf / brief_text also
  import directly from brief_model).

Per-run generator conventions:
- Exactly one output-path line of the form `out = r"..."` (the health-checkup
  baseline harness redirects that line to run scripts safely).
- Section items should be `SectionItem(company, category, dateline, links,
  summary, assessment)` named records (see `brief_model`). The builder also
  accepts the legacy positional 7-tuple (company, category, dateline, links,
  summary, assessment, draft) and normalizes it — the trailing `draft` slot is
  inert (retired 2026-06-18) and is simply ignored. Pass `SectionItem` for new
  generators; legacy tuples remain fully supported so historical fixtures still
  run unchanged.
- The `assessment` slot (the editorial "Assessment:" line) was retired 2026-06-15
  and is no longer rendered. Pass None (or omit from SectionItem).
- exec_items entries are `(display_name, one_line_summary)` or, for lead names
  that live ONLY in the Executive Summary, `(display_name, one_line_summary,
  links)` where `links` is a list of (label, url) rendered inline so promoted
  items stay traceable.
- The single-coverage rule (principal, 2026-06-18): a lead name appears in the
  Executive Summary OR a sector section, never both. Generators MUST call
  `assert_single_coverage(exec_items, sections)` before building (enforced at
  the content layer, not inside the renderers, so the renderers can still
  reproduce historical double-covered briefs for the baseline).
"""
import os

from brief_model import (  # re-exported so existing callers/tests keep importing from brief_docx
    _TICKER_RE, _star_keys, _make_star_test, _is_pseudo_entry, assert_single_coverage,
    build_blocks,
    TitleBlock, Heading1, ExecItem, CompanyItem, Bullet, Paragraph,
    PlainBullet, NumberedItem, Note)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
RED = RGBColor(0xFF, 0x00, 0x00)
STAR = "★"  # red-star mark on the 1-2 most postable stories (sole permitted color accent)
FONT = "Times New Roman"


def add_hyperlink(paragraph, text, url):
    """Black, underlined hyperlink (functional but visually restrained)."""
    if not text or not url:
        raise ValueError("add_hyperlink requires non-empty text and url "
                         "(an empty label renders as an invisible link)")
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), FONT); fonts.set(qn("w:hAnsi"), FONT); rPr.append(fonts)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "000000"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number_footer(doc):
    """Confidentiality line + centered PAGE field. Call once per document —
    a second call appends a duplicate footer block."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Privileged & Confidential — For Internal Distribution")
    r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = BLACK
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run()
    run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = BLACK
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)



# ---------------------------------------------------------------------------
# Block painters - one per block type; each takes (doc, block) and returns None.
# ---------------------------------------------------------------------------

def _paint_title(doc, b):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PUBCO BRIEF"); r.bold = True; r.font.size = Pt(16)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(b.date_line); r.font.size = Pt(11)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRIVILEGED & CONFIDENTIAL — FOR INTERNAL DISTRIBUTION"); r.bold = True; r.font.size = Pt(9)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(b.coverage_note)
    r.italic = True; r.font.size = Pt(10)


def _paint_h1(doc, b):
    doc.add_heading(b.text, 1)


def _paint_exec(doc, b):
    p = doc.add_paragraph(style="List Number")
    if b.starred:
        r = p.add_run(STAR + " ")
        r.bold = True
        r.font.color.rgb = RED
    r = p.add_run(b.name + ". ")
    r.bold = True
    p.add_run(b.summary)
    if b.links:
        p.add_run(" [")
        for i, (label, url) in enumerate(b.links):
            if i:
                p.add_run("; ")
            add_hyperlink(p, label, url)
        p.add_run("]")


def _paint_company(doc, b):
    h = doc.add_heading("", 2)
    if b.starred:
        r = h.add_run(STAR + " ")
        r.font.color.rgb = RED
    h.add_run(b.company)
    p = doc.add_paragraph()
    p.add_run(b.category + " — " + b.dateline)
    if b.links:
        p.add_run(" — Source: ")
        for i, (label, url) in enumerate(b.links):
            if i:
                p.add_run("; ")
            add_hyperlink(p, label, url)
        p.add_run(".")
    body = doc.add_paragraph(b.summary)
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if b.assessment:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run("Assessment: "); r.bold = True
        p.add_run(b.assessment)


def _paint_bullet(doc, b):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(b.title + " — "); r.bold = True
    p.add_run(b.text + " [")
    add_hyperlink(p, b.src_label, b.src_url)
    p.add_run("]")


def _paint_paragraph(doc, b):
    p = doc.add_paragraph(b.text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _paint_plain_bullet(doc, b):
    doc.add_paragraph(b.text, style="List Bullet")


def _paint_numbered(doc, b):
    doc.add_paragraph(b.text, style="List Number")


def _paint_note(doc, b):
    intro = doc.add_paragraph()
    r = intro.add_run(b.text)
    r.italic = True; r.font.size = Pt(9)


_PAINT = {
    TitleBlock:   _paint_title,
    Heading1:     _paint_h1,
    ExecItem:     _paint_exec,
    CompanyItem:  _paint_company,
    Bullet:       _paint_bullet,
    Paragraph:    _paint_paragraph,
    PlainBullet:  _paint_plain_bullet,
    NumberedItem: _paint_numbered,
    Note:         _paint_note,
}


def render_docx(blocks):
    """Paint a recipe (list of block records) into a Document and return it."""
    doc = Document()
    apply_formal_styles(doc)
    add_page_number_footer(doc)
    for b in blocks:
        _PAINT[type(b)](doc, b)
    return doc

def render_brief(*, date_line, coverage_note, exec_items, sections, starred=(),
                 adjacent=(), adjacent_note=None, market_notes=(), trackers=(),
                 catchup=None, hygiene=()):
    """Build and return the complete formal brief as a Document.

    Thin Word-backend entry point: builds the recipe (brief_model.build_blocks,
    which encodes the June 12, 2026 format and runs the content rules once) and
    paints it via render_docx. Per-run scripts supply content only:
    - exec_items: [(display_name, one_line_summary)] or, for lead names that
                  now live ONLY in the exec summary, (display_name, summary,
                  links) where links is [(label, url)] rendered inline.
    - sections:   [(sector, [SectionItem records — or legacy 7-tuples; see the
                  module docstring])]. `assessment` renders when present
                  (historical briefs); the retired `draft` slot is ignored.
    - starred:    display names or bare tickers for the 1-2 ★ items
    - adjacent:   [(title, text, source_label, source_url)] bullets, or pass
                  adjacent_note (str) for a no-transactions paragraph instead
    - market_notes: [(title, text, source_label, source_url)] — the "Market
                  Notes" section: broad-market conversation-starters beyond the
                  watchlist (IPOs, mega-deals, macro). Same shape as adjacent.
    - trackers:   [str] story-tracker bullet lines
    - catchup:    optional (heading, body) pre-window catch-up section; body
                  may be a str (single justified paragraph) or a list of
                  one-liners (rendered as bullets — the preferred form)
    - hygiene:    [str] numbered watchlist-hygiene appendix lines

    The "Companies with No Reportable Developments" (quiet) and "Run Notes"
    sections were retired 2026-06-15 (SKILL.md Step 9 bans them from the
    deliverable); their parameters were removed so a generator can no longer
    resurrect them — passing quiet=/run_notes= now raises TypeError.
    """
    return render_docx(build_blocks(
        date_line=date_line, coverage_note=coverage_note, exec_items=exec_items,
        sections=sections, starred=starred, adjacent=adjacent,
        adjacent_note=adjacent_note, market_notes=market_notes, trackers=trackers,
        catchup=catchup, hygiene=hygiene))


def save_brief(doc, out):
    """Save the brief, refusing to overwrite an existing file.

    Versioning is the caller's job (Step 11: suffix ' v2', ' v3' before the
    extension when today's file exists). This guard exists so rerunning an
    old generator script can never destroy a shipped deliverable."""
    if os.path.exists(out):
        raise FileExistsError(
            "refusing to overwrite existing brief: %s — pick the next "
            "' v2'/' v3' name per SKILL.md Step 11" % out)
    doc.save(out)


def apply_formal_styles(doc):
    """0.5" margins; Normal + Heading 1/2 forced to black Times New Roman."""
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(0.5))

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK

    for h in ("Heading 1", "Heading 2"):
        s = doc.styles[h]
        s.font.name = FONT
        s.font.color.rgb = BLACK
        s.font.bold = True
        s.font.size = Pt(13) if h == "Heading 1" else Pt(11.5)
        # Word links heading fonts to the theme via rFonts asciiTheme; clear it so TNR sticks
        rpr = s.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            for a in ("w:asciiTheme", "w:hAnsiTheme"):
                if rfonts.get(qn(a)) is not None:
                    del rfonts.attrib[qn(a)]
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)
